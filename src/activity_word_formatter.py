"""
Activity Word Formatter (Teacher-Friendly Version)
Generates professional Word documents with visual analytics for teachers.
Hides technical NLP details; shows only actionable insights.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging
from typing import Dict, List

logger = logging.getLogger(__name__)


def create_activity_report(analysis_results: Dict) -> Document:
    """
    Create a teacher-friendly Word document from activity analysis results.
    
    Structure:
    - Page 1: Visual Analytics Dashboard
    - Page 2: Executive Summary
    - Page 3: Q1 - Top Learning Responses
    - Page 4: Q2 - Top Student Questions
    - Page 5: Q3 - Top Engagement Reflections
    
    Args:
        analysis_results: Dictionary containing complete analysis results
        
    Returns:
        Document: python-docx Document object
    """
    doc = Document()
    metadata = analysis_results['metadata']
    q1 = analysis_results['q1_analysis']
    q2 = analysis_results['q2_analysis']
    q3 = analysis_results['q3_analysis']
    recommendations = analysis_results.get('recommendations', [])
    
    # Page 1: Visual Analytics Dashboard
    _add_visual_dashboard(doc, analysis_results)
    doc.add_page_break()
    
    # Page 2: Executive Summary
    _add_executive_summary(doc, analysis_results, recommendations)
    doc.add_page_break()
    
    # Page 3: Q1 - Top Learning Responses
    _add_q1_section(doc, q1)
    doc.add_page_break()
    
    # Page 4: Q2 - Top Student Questions
    _add_q2_section(doc, q2)
    doc.add_page_break()
    
    # Page 5: Q3 - Top Engagement Reflections
    _add_q3_section(doc, q3)
    
    logger.info("Teacher-friendly activity report created successfully")
    return doc


def _add_visual_dashboard(doc: Document, results: Dict):
    """Create visual analytics dashboard with tables."""
    metadata = results['metadata']
    q1 = results['q1_analysis']
    q2 = results['q2_analysis']
    q3 = results['q3_analysis']
    summary = results['summary']
    
    # Title
    title = doc.add_heading('Activity Learning Analytics Dashboard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Generation info
    p = doc.add_paragraph()
    p.add_run('Report Generated: ').bold = True
    p.add_run(metadata['timestamp'][:19].replace('T', ' '))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========================================
    # Table 1: Overall Class Engagement
    # ========================================
    doc.add_heading('Overall Class Engagement', 2)
    
    table1 = doc.add_table(rows=5, cols=3)
    table1.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table1.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Count'
    header_cells[2].text = 'Percentage'
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    total_students = metadata['total_students']
    
    engagement_data = [
        ('Total Students', str(total_students), '100%'),
        ('Q1 Learning Summaries', str(summary['q1_responses_analyzed']), 
         f"{summary['q1_responses_analyzed']/total_students*100:.0f}%" if total_students > 0 else '0%'),
        ('Q2 Questions Submitted', str(summary['q2_responses_analyzed']),
         f"{summary['q2_responses_analyzed']/total_students*100:.0f}%" if total_students > 0 else '0%'),
        ('Q3 Reflections Shared', str(summary['q3_responses_analyzed']),
         f"{summary['q3_responses_analyzed']/total_students*100:.0f}%" if total_students > 0 else '0%'),
    ]
    
    for i, (metric, count, pct) in enumerate(engagement_data, 1):
        row = table1.rows[i].cells
        row[0].text = metric
        row[1].text = count
        row[2].text = pct
    
    doc.add_paragraph()
    
    # ========================================
    # Table 2: Q1 Learning Outcomes (Cognitive Domain)
    # ========================================
    doc.add_heading('Q1: Learning Outcomes (Cognitive Domain)', 2)
    
    cognitive = q1.get('cognitive_categorization', {})
    
    table2 = doc.add_table(rows=3, cols=4)
    table2.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table2.rows[0].cells
    header_cells[0].text = 'Category'
    header_cells[1].text = 'Students'
    header_cells[2].text = 'Percentage'
    header_cells[3].text = 'Description'
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    # Learned Well row
    learned_well = cognitive.get('learned_well', {})
    row1 = table2.rows[1].cells
    row1[0].text = '✓ Learned Well'
    row1[1].text = str(learned_well.get('count', 0))
    row1[2].text = f"{learned_well.get('percentage', 0)}%"
    row1[3].text = learned_well.get('description', 'Clear understanding demonstrated')
    
    # Make "Learned Well" green
    for para in row1[0].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 128, 0)
    
    # Needs Reinforcement row
    needs_reinforcement = cognitive.get('needs_reinforcement', {})
    row2 = table2.rows[2].cells
    row2[0].text = '⚠ Needs Reinforcement'
    row2[1].text = str(needs_reinforcement.get('count', 0))
    row2[2].text = f"{needs_reinforcement.get('percentage', 0)}%"
    row2[3].text = needs_reinforcement.get('description', 'May need additional support')
    
    # Make "Needs Reinforcement" orange
    for para in row2[0].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 140, 0)
    
    doc.add_paragraph()
    
    # ========================================
    # Table 3: Q3 Student Engagement (Affective Domain)
    # ========================================
    doc.add_heading('Q3: Student Engagement (Affective Domain)', 2)
    
    affective = q3.get('affective_categorization', {})
    
    table3 = doc.add_table(rows=3, cols=4)
    table3.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table3.rows[0].cells
    header_cells[0].text = 'Category'
    header_cells[1].text = 'Students'
    header_cells[2].text = 'Percentage'
    header_cells[3].text = 'Description'
    for cell in header_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    
    # Wants to Explore row
    wants_to_explore = affective.get('wants_to_explore', {})
    row1 = table3.rows[1].cells
    row1[0].text = '🔍 Wants to Explore Further'
    row1[1].text = str(wants_to_explore.get('count', 0))
    row1[2].text = f"{wants_to_explore.get('percentage', 0)}%"
    row1[3].text = wants_to_explore.get('description', 'Shows curiosity and exploration intent')
    
    # Make "Wants to Explore" blue
    for para in row1[0].paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 100, 200)
    
    # General Interest row
    general_interest = affective.get('general_interest', {})
    row2 = table3.rows[2].cells
    row2[0].text = '💭 General Interest'
    row2[1].text = str(general_interest.get('count', 0))
    row2[2].text = f"{general_interest.get('percentage', 0)}%"
    row2[3].text = general_interest.get('description', 'Expressed interest but no specific direction')


def _add_executive_summary(doc: Document, results: Dict, recommendations: List[str]):
    """Add executive summary with key findings."""
    metadata = results['metadata']
    q1 = results['q1_analysis']
    q3 = results['q3_analysis']
    
    doc.add_heading('Executive Summary', 1)
    
    # Activity Context
    doc.add_heading('Activity Context', 2)
    doc.add_paragraph(metadata['activity_description'])
    
    doc.add_paragraph()
    
    # Key Findings
    doc.add_heading('Key Findings', 2)
    
    cognitive = q1.get('cognitive_categorization', {})
    affective = q3.get('affective_categorization', {})
    
    findings = [
        f"• {cognitive.get('learned_well', {}).get('count', 0)} students ({cognitive.get('learned_well', {}).get('percentage', 0)}%) demonstrated strong learning",
        f"• {len(results['q2_analysis'].get('top_10_questions', []))} thoughtful questions identified for follow-up",
        f"• {affective.get('wants_to_explore', {}).get('count', 0)} students ({affective.get('wants_to_explore', {}).get('percentage', 0)}%) want to explore the topic further"
    ]
    
    for finding in findings:
        doc.add_paragraph(finding)
    
    doc.add_paragraph()
    
    # Recommended Actions
    doc.add_heading('Recommended Actions', 2)
    
    if recommendations:
        for i, rec in enumerate(recommendations, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. ").bold = True
            p.add_run(rec)
    else:
        doc.add_paragraph("Analysis completed successfully. Review top responses below.")
    
    doc.add_paragraph()
    
    # Quick Stats Box
    doc.add_heading('At a Glance', 2)
    
    total = metadata['total_students']
    learned_count = cognitive.get('learned_well', {}).get('count', 0)
    explore_count = affective.get('wants_to_explore', {}).get('count', 0)
    
    stats_table = doc.add_table(rows=1, cols=3)
    stats_table.style = 'Light Grid Accent 5'
    
    cell1 = stats_table.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.add_run(f'{total}').bold = True
    p1.add_run('\nTotal Students')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cell2 = stats_table.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.add_run(f'{learned_count}').bold = True
    p2.add_run('\nLearned Well')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cell3 = stats_table.rows[0].cells[2]
    p3 = cell3.paragraphs[0]
    p3.add_run(f'{explore_count}').bold = True
    p3.add_run('\nWant to Explore')
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_q1_section(doc: Document, q1_results: Dict):
    """Q1 section - Shows concept keywords and top learning responses."""
    doc.add_heading('Q1: What Students Learned', 1)
    
    doc.add_paragraph(
        'Question: Write three things you learned during this lesson.'
    )
    
    # Display concept keywords extracted from activity (cognitive domain measurement)
    concept_keywords = q1_results.get('concept_keywords', [])
    if concept_keywords:
        doc.add_heading('Keywords: (Concepts Taught - Cognitive Domain)', 2)
        p = doc.add_paragraph('Example - ')
        keywords_text = ', '.join(concept_keywords[:15])  # Show top 15 concepts
        p.add_run(keywords_text)
        doc.add_paragraph()
    
    cognitive = q1_results.get('cognitive_categorization', {})
    learned_count = cognitive.get('learned_well', {}).get('count', 0)
    total = q1_results.get('total_analyzed', 0)
    
    p = doc.add_paragraph()
    p.add_run('Summary: ').bold = True
    p.add_run(f'{learned_count} out of {total} students demonstrated strong learning.')
    
    doc.add_paragraph()
    
    top_responses = q1_results.get('top_10_responses', [])
    
    doc.add_heading(f'Top {len(top_responses)} Learning Responses', 2)
    
    if not top_responses:
        doc.add_paragraph('No responses to display.')
        return
    
    for i, resp in enumerate(top_responses, 1):
        # Simple heading: just number and student ID
        doc.add_heading(f'{i}. Student {resp.get("student_id", "Unknown")}', 3)
        
        # Response text - clean and simple
        p = doc.add_paragraph()
        response_text = resp.get('response', 'No response')
        p.add_run(f'"{response_text}"')
        p.style = 'Quote'
        
        doc.add_paragraph()  # Spacing


def _add_q2_section(doc: Document, q2_results: Dict):
    """Q2 section - Shows concept-based themes with example questions and frequency."""
    doc.add_heading('Q2: Student Questions', 1)
    
    doc.add_paragraph(
        'Question: Write two questions you have about the instructional material discussed in the lesson.'
    )
    
    total = q2_results.get('total_analyzed', 0)
    
    p = doc.add_paragraph()
    p.add_run('Summary: ').bold = True
    p.add_run(f'{total} questions collected from students.')
    
    doc.add_paragraph()
    
    # Display theme-based grouping with frequency
    themes = q2_results.get('themes', {})
    if themes:
        doc.add_heading('Questions by Concept Theme (FAQ - Top Categories)', 2)
        
        # Create table: Theme | Example Student Questions | Frequency
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Theme'
        header_cells[1].text = 'Example Student Questions'
        header_cells[2].text = 'Frequency'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add theme rows
        for theme_name, theme_data in list(themes.items())[:10]:  # Top 10 themes
            questions = theme_data.get('questions', [])
            frequency = theme_data.get('frequency', 0)
            
            if questions:
                # Get first example question
                example_question = questions[0].get('question', 'No question')
                if len(example_question) > 150:
                    example_question = example_question[:150] + '...'
                
                row_cells = table.add_row().cells
                row_cells[0].text = theme_name
                row_cells[1].text = example_question
                row_cells[2].text = str(frequency)
        
        doc.add_paragraph()
    
    # Also show top 10 questions
    top_questions = q2_results.get('top_10_questions', [])
    
    if top_questions:
        doc.add_heading(f'Top {len(top_questions)} Student Questions', 2)
        
        for i, quest in enumerate(top_questions, 1):
            # Simple heading: just number and student ID
            doc.add_heading(f'{i}. Student {quest.get("student_id", "Unknown")}', 3)
            
            # Question text - clean and simple
            p = doc.add_paragraph()
            question_text = quest.get('question', 'No question')
            p.add_run(f'"{question_text}"')
            p.style = 'Quote'
            
            doc.add_paragraph()  # Spacing


def _add_q3_section(doc: Document, q3_results: Dict):
    """Q3 section - Shows content-related vs pedagogy-related themes."""
    doc.add_heading('Q3: Student Interests & Exploration', 1)
    
    doc.add_paragraph(
        'Question: Write which one aspect you found most interesting or something you would like to explore further related to the topic discussed.'
    )
    
    affective = q3_results.get('affective_categorization', {})
    explore_count = affective.get('wants_to_explore', {}).get('count', 0)
    total = q3_results.get('total_analyzed', 0)
    
    p = doc.add_paragraph()
    p.add_run('Summary: ').bold = True
    p.add_run(f'{explore_count} out of {total} students showed interest in further exploration.')
    
    doc.add_paragraph()
    
    # Display content-related themes
    content_themes = q3_results.get('content_themes', {})
    pedagogy_themes = q3_results.get('pedagogy_themes', {})
    
    if content_themes or pedagogy_themes:
        doc.add_heading('1 Aspect They Found Interesting', 2)
        
        # Create table: Concept/Theme | Example Phrasing
        if content_themes or pedagogy_themes:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Concept/Theme'
            header_cells[1].text = 'Example Phrasing'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add content-related themes
            for theme_name, theme_data in content_themes.items():
                example = theme_data.get('example_phrasing', '')
                if len(example) > 200:
                    example = example[:200] + '...'
                
                row_cells = table.add_row().cells
                row_cells[0].text = theme_name
                row_cells[1].text = example
            
            # Add pedagogy-related themes (with visual distinction)
            for theme_name, theme_data in pedagogy_themes.items():
                example = theme_data.get('example_phrasing', '')
                if len(example) > 200:
                    example = example[:200] + '...'
                
                row_cells = table.add_row().cells
                # Mark pedagogy themes (could use formatting if needed)
                row_cells[0].text = theme_name
                row_cells[1].text = example
            
            doc.add_paragraph()
            
            # Add categorization labels
            if content_themes and pedagogy_themes:
                p = doc.add_paragraph()
                p.add_run('Content-related: ').bold = True
                p.add_run('Themes focusing on concepts, topics, and subject matter.')
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('Pedagogy-related: ').bold = True
                p.add_run('Themes focusing on teaching methods, activities, and learning approaches.')
                doc.add_paragraph()
    
    # Also show top 10 responses
    top_responses = q3_results.get('top_10_responses', [])
    
    if top_responses:
        doc.add_heading(f'Top {len(top_responses)} Student Insights', 2)
        
        for i, resp in enumerate(top_responses, 1):
            # Simple heading: just number and student ID
            doc.add_heading(f'{i}. Student {resp.get("student_id", "Unknown")}', 3)
            
            # Response text - clean and simple
            p = doc.add_paragraph()
            response_text = resp.get('response', 'No response')
            p.add_run(f'"{response_text}"')
            p.style = 'Quote'
            
            doc.add_paragraph()  # Spacing
