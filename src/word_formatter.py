"""
Word Document Formatter for Generated Questions
Converts question data to formatted Word documents.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging
from typing import Dict, List

logger = logging.getLogger(__name__)


class WordFormatter:
    """Formats generated questions into Word documents."""
    
    def __init__(self):
        """Initialize the Word formatter."""
        pass
    
    def create_document(self, concept_data: Dict) -> Document:
        """
        Create a formatted Word document from concept data.
        
        Args:
            concept_data: Dictionary containing concept information and questions
            
        Returns:
            Document: python-docx Document object
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Practice Questions: {concept_data["concept_name"]}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        doc.add_paragraph()
        self._add_metadata_section(doc, concept_data)
        
        # Add page break
        doc.add_page_break()
        
        # Add questions by difficulty level
        levels = concept_data.get('levels', {})
        
        if 'levels' in levels:
            # Handle nested structure
            levels = levels['levels']
        
        for difficulty in ['beginner', 'intermediate', 'advanced']:
            if difficulty in levels and 'questions' in levels[difficulty]:
                questions = levels[difficulty]['questions']
                if questions:
                    self._add_difficulty_section(doc, difficulty, questions)
                    doc.add_page_break()
        
        # Add answer key at the end
        doc.add_page_break()
        self._add_answer_key(doc, levels)
        
        return doc
    
    def _add_metadata_section(self, doc: Document, concept_data: Dict):
        """Add metadata section to document."""
        # Concept information
        p = doc.add_paragraph()
        p.add_run('Concept: ').bold = True
        p.add_run(concept_data.get('concept_name', 'N/A'))
        
        # Category
        p = doc.add_paragraph()
        p.add_run('Category: ').bold = True
        p.add_run(concept_data.get('category', 'N/A').title())
        
        # Programming language (if applicable)
        if concept_data.get('language'):
            p = doc.add_paragraph()
            p.add_run('Language: ').bold = True
            p.add_run(concept_data.get('language', 'N/A').upper())
        
        # Affected students
        students = concept_data.get('affected_students', [])
        p = doc.add_paragraph()
        p.add_run('Students Affected: ').bold = True
        p.add_run(str(len(students)))
        
        # Instructions
        doc.add_paragraph()
        instructions = doc.add_paragraph()
        instructions.add_run('Instructions:').bold = True
        doc.add_paragraph('• Read each question carefully')
        doc.add_paragraph('• Select the best answer from the options provided')
        doc.add_paragraph('• Review the explanations after completing all questions')
        doc.add_paragraph('• Use hints if you get stuck')
    
    def _add_difficulty_section(self, doc: Document, difficulty: str, questions: List[Dict]):
        """Add a section for a specific difficulty level."""
        # Skip if no questions
        if not questions or len(questions) == 0:
            return
        
        # Section heading
        heading = doc.add_heading(f'{difficulty.title()} Level Questions', 1)
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)
        
        # Add each question
        for idx, question in enumerate(questions, 1):
            try:
                self._add_question(doc, idx, question, difficulty)
                doc.add_paragraph()  # Add spacing between questions
            except Exception as e:
                logger.error(f"Error adding question {idx} for {difficulty}: {e}")
                continue
    
    def _add_question(self, doc: Document, number: int, question_data: Dict, difficulty: str):
        """Add a single question to the document."""
        # Question number and text
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f'Question {number}: ')
        q_run.bold = True
        q_run.font.size = Pt(12)

        # Get question text (or title for advanced)
        q_text_content = question_data.get('question') or question_data.get('title') or 'No question text'
        q_text = q_para.add_run(q_text_content)
        q_text.font.size = Pt(11)

        # Add description (for advanced questions - separate from title)
        if question_data.get('description') and question_data.get('title'):
            desc_para = doc.add_paragraph(question_data['description'])
            desc_para.paragraph_format.left_indent = Inches(0.25)

        # Add function signature (advanced)
        if question_data.get('function_signature'):
            sig_para = doc.add_paragraph(question_data['function_signature'])
            sig_para.style = 'Intense Quote'
            sig_para.paragraph_format.left_indent = Inches(0.5)

        # Add constraints (advanced)
        if question_data.get('constraints'):
            const_para = doc.add_paragraph()
            const_para.add_run('Constraints: ').bold = True
            for constraint in question_data['constraints']:
                c_para = doc.add_paragraph(f'• {constraint}')
                c_para.paragraph_format.left_indent = Inches(0.5)

        # Add code if present
        if 'code' in question_data and question_data['code']:
            code_para = doc.add_paragraph(question_data['code'])
            code_para.style = 'Intense Quote'
            code_para.paragraph_format.left_indent = Inches(0.5)

        # Add options
        options = question_data.get('options', {})
        if options and isinstance(options, dict):
            try:
                for option_key in sorted(options.keys()):
                    option_text = options.get(option_key, '')
                    option_para = doc.add_paragraph(f'{option_key}. {option_text}')
                    option_para.paragraph_format.left_indent = Inches(0.5)
                    option_para.style = 'List Bullet'
            except Exception as e:
                logger.error(f"Error adding options: {e}")

        # Add hints (from question_data directly for advanced)
        hints = question_data.get('hints', [])
        if not hints:
            feedback = question_data.get('feedback', {})
            hints = feedback.get('hints', [])

        if hints:
            hint_para = doc.add_paragraph()
            hint_run = hint_para.add_run('💡 Hints: ')
            hint_run.italic = True
            hint_run.font.size = Pt(10)

            for hint in hints:
                hint_p = doc.add_paragraph(f'• {hint}')
                hint_p.paragraph_format.left_indent = Inches(0.75)
                hint_p.runs[0].font.size = Pt(9)
                hint_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_answer_key(self, doc: Document, levels: Dict):
        """Add answer key section at the end."""
        doc.add_heading('Answer Key & Explanations', 1)
        
        for difficulty in ['beginner', 'intermediate', 'advanced']:
            if difficulty in levels and 'questions' in levels[difficulty]:
                questions = levels[difficulty]['questions']
                if not questions or len(questions) == 0:
                    continue
                
                # Difficulty subsection
                doc.add_heading(f'{difficulty.title()} Level', 2)
                
                for idx, question in enumerate(questions, 1):
                    try:
                        self._add_answer_explanation(doc, idx, question)
                    except Exception as e:
                        logger.error(f"Error adding answer for question {idx} in {difficulty}: {e}")
                        continue
    
    def _add_answer_explanation(self, doc: Document, number: int, question_data: Dict):
        """Add answer and explanation for a question."""
        try:
            # Question number and answer
            ans_para = doc.add_paragraph()
            ans_run = ans_para.add_run(f'Question {number}: ')
            ans_run.bold = True

            # Get answer (correct_answer or solution_approach for advanced)
            answer = (question_data.get('correct_answer') or
                     question_data.get('solution') or
                     'See solution approach below')
            answer_run = ans_para.add_run(f'Answer: {answer}')
            answer_run.font.color.rgb = RGBColor(0, 128, 0)
            answer_run.bold = True

            # Explanation/Solution approach
            explanation = (question_data.get('explanation') or
                          question_data.get('solution_approach') or
                          question_data.get('feedback', {}).get('general_explanation', 'No explanation available'))
            exp_para = doc.add_paragraph(str(explanation))
            exp_para.paragraph_format.left_indent = Inches(0.5)
            if exp_para.runs and len(exp_para.runs) > 0:
                exp_para.runs[0].font.size = Pt(10)

            # Add test cases (for advanced questions)
            test_cases = question_data.get('test_cases', [])
            if test_cases and isinstance(test_cases, list):
                tc_heading = doc.add_paragraph()
                tc_heading.add_run('Test Cases:').bold = True
                tc_heading.paragraph_format.left_indent = Inches(0.5)

                for idx, tc in enumerate(test_cases, 1):
                    tc_para = doc.add_paragraph()
                    tc_para.add_run(f'Test Case {idx}: ').bold = True
                    tc_input = tc.get('input', '')
                    tc_output = tc.get('expected_output', '')
                    tc_explain = tc.get('explanation', '')
                    tc_para.add_run(f"Input: {tc_input} → Output: {tc_output}")
                    if tc_explain:
                        tc_para.add_run(f" ({tc_explain})")
                    tc_para.paragraph_format.left_indent = Inches(0.75)
                    if tc_para.runs:
                        tc_para.runs[0].font.size = Pt(9)

            # Option feedback
            feedback = question_data.get('feedback', {})
            option_feedback = feedback.get('option_feedback', {})
            if option_feedback and isinstance(option_feedback, dict):
                doc.add_paragraph('Detailed feedback:', style='List Bullet')
                for option, feedback_text in sorted(option_feedback.items()):
                    feedback_para = doc.add_paragraph(f'{option}: {feedback_text}')
                    feedback_para.paragraph_format.left_indent = Inches(0.75)
                    if feedback_para.runs and len(feedback_para.runs) > 0:
                        feedback_para.runs[0].font.size = Pt(9)
            
            doc.add_paragraph()  # Add spacing
        except Exception as e:
            logger.error(f"Error in _add_answer_explanation: {e}")
            # Add a simple error message to the document
            doc.add_paragraph(f"[Error displaying answer for Question {number}]")


def create_word_document(concept_data: Dict) -> Document:
    """
    Create a Word document from concept question data.
    
    Args:
        concept_data: Dictionary containing concept information and questions
        
    Returns:
        Document: python-docx Document object
    """
    formatter = WordFormatter()
    return formatter.create_document(concept_data)

